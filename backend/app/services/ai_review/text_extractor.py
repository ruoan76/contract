# -*- coding: utf-8 -*-
"""合同文本提取 — 支持 PDF / DOCX / TXT"""
from __future__ import annotations

import asyncio
import logging
from pathlib import Path
from typing import Any, List, Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from pydantic import BaseModel, Field

from app.core.config import settings
from app.services.ai_review.document_json import build_document_json, render_document_text
from app.services.ai_review.ocr_text_utils import indices_needing_ocr

logger = logging.getLogger(__name__)


def _format_pages_with_markers(pages: list[str], *, ocr_used: bool) -> list[str]:
    """OCR 场景为每页添加分页标记，便于人工核对。"""
    if not ocr_used:
        return pages
    formatted: list[str] = []
    for idx, page in enumerate(pages):
        body = page.strip()
        marker = f"--- 第 {idx + 1} 页 ---"
        formatted.append(f"{marker}\n\n{body}" if body else marker)
    return formatted


def _join_page_text(pages: list[str], *, ocr_used: bool) -> str:
    chunks = _format_pages_with_markers(pages, ocr_used=ocr_used)
    return "\n\n".join(chunks)


# ---------------------------------------------------------------------------
# 数据模型
# ---------------------------------------------------------------------------

class ExtractedText(BaseModel):
    """提取后的合同文本"""

    full_text: str = Field(default="", description="完整纯文本")
    pages: List[str] = Field(default_factory=list, description="每页文本列表")
    tables: List[List[List[str]]] = Field(default_factory=list, description="表格数据")
    metadata: dict[str, Any] = Field(default_factory=dict, description="文件元数据")


# ---------------------------------------------------------------------------
# 公共入口
# ---------------------------------------------------------------------------

async def extract_text(file_path: str, file_type: str) -> ExtractedText:
    """从合同文件中提取文本

    Args:
        file_path: 文件绝对路径
        file_type: 文件类型 (pdf / docx / txt)

    Returns:
        ExtractedText 包含提取结果

    Raises:
        FileNotFoundError: 文件不存在
        ValueError: 不支持的文件类型或文件过大
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    file_size = path.stat().st_size
    if file_size > settings.MAX_FILE_SIZE:
        raise ValueError(
            f"文件大小 {file_size} 字节超过限制 {settings.MAX_FILE_SIZE}"
        )

    file_type = file_type.lower().lstrip(".")

    extractor = _ROUTING.get(file_type)
    if extractor is None:
        raise ValueError(f"不支持的文件类型: {file_type}")

    return await extractor(path, file_type)


# ---------------------------------------------------------------------------
# PDF 提取
# ---------------------------------------------------------------------------

async def _extract_pdf(path: Path, file_type: str) -> ExtractedText:
    """使用 PyMuPDF 提取 PDF 文本，逐页混合 OCR。"""
    loop = asyncio.get_event_loop()
    tables: list[list[list[str]]] = []

    def _read() -> tuple[list[str], list[list[list[str]]], dict[str, Any], int]:
        doc = fitz.open(str(path))
        try:
            metadata = dict(doc.metadata or {})
            page_texts: list[str] = []
            all_tables: list[list[list[str]]] = []

            for page_idx in range(len(doc)):
                page = doc[page_idx]
                # sort=True 改善复杂版式阅读顺序
                text = page.get_text("text", sort=True)
                page_texts.append(text)

            return page_texts, all_tables, metadata, len(doc)
        finally:
            doc.close()

    pages, tables, metadata, raw_page_count = await loop.run_in_executor(None, _read)

    full_text = _join_page_text(pages, ocr_used=False)
    ocr_used = False
    ocr_page_count = 0
    ocr_page_indices: list[int] = []
    ocr_pages_raw: list[str] = []
    ocr_page_meta: list[dict] = []

    if settings.AI_OCR_ENABLED:
        from app.services.ai_review.ocr import ocr_pdf_page_indices, ocr_pdf_pages

        stripped_len = len(full_text.strip())
        if stripped_len < settings.AI_OCR_MIN_CHARS:
            logger.info(
                "PDF 可提取文字过少（%s 字），全量 OCR: %s",
                stripped_len,
                path,
            )

            def _run_full_ocr() -> list[str]:
                raw_buf: list[str] = []
                meta_buf: list[dict] = []
                result = ocr_pdf_pages(
                    path,
                    max_pages=settings.AI_OCR_MAX_PAGES,
                    raw_pages_out=raw_buf,
                    page_meta_out=meta_buf,
                )
                ocr_pages_raw.extend(raw_buf)
                ocr_page_meta.extend(meta_buf)
                return result

            ocr_pages = await loop.run_in_executor(None, _run_full_ocr)
            pages = ocr_pages
            full_text = _join_page_text(pages, ocr_used=True)
            ocr_used = True
            ocr_page_count = len(ocr_pages)
            ocr_page_indices = list(range(len(ocr_pages)))
        else:
            need_indices = indices_needing_ocr(
                pages,
                min_chars=settings.AI_OCR_PAGE_MIN_CHARS,
                gibberish_threshold=settings.AI_OCR_GIBBERISH_RATIO,
            )
            if need_indices:
                logger.info(
                    "PDF 逐页混合 OCR：%s/%s 页需 OCR: %s",
                    len(need_indices),
                    len(pages),
                    path,
                )

                def _run_partial_ocr() -> tuple[list[str], list[str], list[dict]]:
                    raw_buf: list[str] = []
                    meta_buf: list[dict] = []
                    result = ocr_pdf_page_indices(
                        path,
                        need_indices,
                        max_pages=settings.AI_OCR_MAX_PAGES,
                        raw_pages_out=raw_buf,
                        page_meta_out=meta_buf,
                    )
                    return result, raw_buf, meta_buf

                ocr_results, raw_buf, meta_buf = await loop.run_in_executor(None, _run_partial_ocr)
                ocr_page_meta.extend(meta_buf)
                pages_raw_snapshot = list(pages)
                for i, page_index in enumerate(need_indices):
                    pages[page_index] = ocr_results[i]
                    pages_raw_snapshot[page_index] = raw_buf[i]
                ocr_pages_raw = pages_raw_snapshot
                full_text = _join_page_text(pages, ocr_used=True)
                ocr_used = True
                ocr_page_count = len(need_indices)
                ocr_page_indices = need_indices

    full_text_raw = ""
    if ocr_used and ocr_pages_raw:
        full_text_raw = _join_page_text(ocr_pages_raw, ocr_used=True)

    aligned_page_meta: list[dict] = []
    if ocr_used and ocr_page_meta and ocr_page_indices:
        meta_by_idx = dict(zip(ocr_page_indices, ocr_page_meta))
        aligned_page_meta = [meta_by_idx.get(i, {}) for i in range(len(pages))]
    elif pages:
        aligned_page_meta = [{} for _ in pages]

    document = build_document_json(
        pages,
        ocr_page_indices=ocr_page_indices if ocr_used else [],
        ocr_page_meta=aligned_page_meta if ocr_used else None,
        ocr_used=ocr_used,
    )
    full_text = render_document_text(document)

    return ExtractedText(
        full_text=full_text,
        pages=pages,
        tables=tables,
        metadata={
            "file_type": file_type,
            "page_count": len(pages) or raw_page_count,
            "source_metadata": metadata,
            "ocr_used": ocr_used,
            "ocr_page_count": ocr_page_count,
            "ocr_page_indices": ocr_page_indices,
            "full_text_raw": full_text_raw or None,
            "layout_version": "v2",
            "layout_engine": settings.AI_OCR_LAYOUT,
            "ocr_engine": settings.AI_OCR_ENGINE if ocr_used else None,
            "ocr_page_meta": aligned_page_meta if ocr_used else None,
            "ocr_needs_review": any(m.get("needs_review") for m in aligned_page_meta) if ocr_used else False,
            "document_json": document.model_dump(),
        },
    )


# ---------------------------------------------------------------------------
# DOCX 提取
# ---------------------------------------------------------------------------

async def _extract_docx(path: Path, file_type: str) -> ExtractedText:
    """使用 python-docx 提取 DOCX 文本"""
    loop = asyncio.get_event_loop()

    def _read() -> tuple[str, list[str], list[list[list[str]]]]:
        doc = DocxDocument(str(path))

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        tables: list[list[list[str]]] = []
        for table in doc.tables:
            table_data: list[list[str]] = []
            for row in table.rows:
                row_data: list[str] = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)

        full_text = "\n".join(paragraphs)
        return full_text, paragraphs, tables

    full_text, paragraphs, tables = await loop.run_in_executor(None, _read)

    document = build_document_json(
        paragraphs if paragraphs else [full_text],
        ocr_used=False,
        include_page_markers=False,
    )

    return ExtractedText(
        full_text=render_document_text(document) or full_text,
        pages=paragraphs,
        tables=tables,
        metadata={
            "file_type": file_type,
            "para_count": len(paragraphs),
            "layout_version": "v2",
            "document_json": document.model_dump(),
        },
    )


# ---------------------------------------------------------------------------
# TXT 提取
# ---------------------------------------------------------------------------

async def _extract_txt(path: Path, file_type: str) -> ExtractedText:
    """提取纯文本文件（自动处理编码）"""
    loop = asyncio.get_event_loop()

    def _read() -> str:
        for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                return path.read_text(encoding=encoding)
            except (UnicodeDecodeError, ValueError):
                continue
        return path.read_text(encoding="utf-8", errors="replace")

    text = await loop.run_in_executor(None, _read)
    lines = text.splitlines()

    document = build_document_json(
        ["\n".join(lines)] if lines else [text],
        ocr_used=False,
        include_page_markers=False,
    )

    return ExtractedText(
        full_text=render_document_text(document) or text,
        pages=lines,
        tables=[],
        metadata={
            "file_type": file_type,
            "line_count": len(lines),
            "layout_version": "v2",
            "document_json": document.model_dump(),
        },
    )


# ---------------------------------------------------------------------------
# 路由表
# ---------------------------------------------------------------------------

_ROUTING: dict[str, Any] = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "txt": _extract_txt,
}
